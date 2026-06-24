"""
Generate Breadstick Canvas Capability Overview as a DOCX.
Writes to pipeline/docs/Breadstick_Canvas_Overview_<date>.docx, then caller uploads to Drive.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import date
from pathlib import Path

def shade_cell(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tc_pr.append(shd)

def h1(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.space_after = Pt(4)

def h2(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xc9, 0xa2, 0x27)  # gold accent
    p.space_before = Pt(12)
    p.space_after = Pt(2)

def h3(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x2d, 0x50, 0x16)
    p.space_before = Pt(6)

def para(doc, txt, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def bullet(doc, txt, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    # Support simple **bold** inline
    parts = txt.split('**')
    for i, part in enumerate(parts):
        r = p.add_run(part)
        r.font.size = Pt(10)
        if i % 2 == 1:
            r.bold = True
    return p

def code_line(doc, txt):
    p = doc.add_paragraph()
    r = p.add_run(txt)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.3)

# ── Build ───────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

h1(doc, 'Breadstick Canvas — Capability Overview')
para(doc, f'Generated {date.today().isoformat()}', italic=True, size=9)

# ── TL;DR ────────────────────────────────────────────────────────────────
h2(doc, 'TL;DR')
para(doc,
     'The Canvas is already production-ready for four content pipelines (UGC, Carousel Video, 16-Gami, '
     'Frame Sandwich). 23 node types, all fully implemented — no stubs. The FFmpeg Grade node has '
     '9 adjustable sliders and 6 presets. Honest answer on "n8n but more powerful": closer to true than '
     'you think, but in a specific way — Breadstick is domain-built for content, n8n is general. '
     'Three 1-day builds would meaningfully level it up: subtitle burn-in, custom LUT upload, '
     'and a Notion-to-script feeder node.')

# ── Mental Model ─────────────────────────────────────────────────────────
h2(doc, 'Is this really "n8n but more powerful"?')
para(doc,
     'Yes — in one direction. n8n is a general workflow engine: HTTP calls, DB queries, polling, '
     'cron. Breadstick is narrower but deeper — every node is hand-crafted for AI content production. '
     'Where n8n gives you 500 integrations and no opinion, Breadstick gives you 23 nodes that already '
     'know how to talk to kie.ai, Anthropic, Blotato, Remotion, and FFmpeg in the exact shape content '
     'production needs.')
para(doc,
     'The practical difference:')
bullet(doc, '**n8n**: "I need to POST to this URL with a JSON body, poll for status, then fetch the result." You build it.')
bullet(doc, '**Breadstick**: "Generate a Kling 3.0 video from first+last frame with origami unfold motion." The node already does it.')
para(doc,
     'The weakness is the same as the strength: narrow domain. You cannot use Breadstick to sync '
     'HubSpot with Notion. But for AI content, this is the only tool that does it this well without '
     'a paid subscription.')

# ── Canvas at a glance ───────────────────────────────────────────────────
h2(doc, 'The Canvas at a glance — 4 content lanes')

h3(doc, '1. UGC Lane — Scripted short-form influencer video')
code_line(doc, 'Character → Ingredients → UGC-Gen → Clip Splitter → Clip Frames + Avatar → UGC Video → FFmpeg Grade → Blotato')
para(doc, 'Output: 5s–55s influencer clips, color-graded, posted to IG/TikTok/FB/X.', size=9)

h3(doc, '2. Carousel Video Lane — Branded slides with embedded motion')
code_line(doc, 'Niche Script Gen → 16-GAMI Art → Carousel → Video Prompt → KIE Img2Vid → Remotion Compositor → FFmpeg Grade')
para(doc, 'Output: Skyframe-branded carousel slides with motion video embedded in art zones.', size=9)

h3(doc, '3. 16-Gami Lane — Origami art slide galleries')
code_line(doc, 'Niche Script Gen → 16-GAMI Art → Carousel → (static slides)')
para(doc, 'Output: Multi-slide origami image sets for static carousel posts.', size=9)

h3(doc, '4. Frame Sandwich Lane — Paper-motion micro-videos')
code_line(doc, 'Title Card (first frame) + 16-GAMI (last frame) → Frame Sandwich → FFmpeg Grade')
para(doc, 'Output: 3s–10s stop-motion-style paper videos (Kling 3.0 first+last frame model).', size=9)

# ── FFmpeg Node ──────────────────────────────────────────────────────────
h2(doc, 'The FFmpeg Grade Node — full capability today')
para(doc, 'Location: server.js lines 252–368. Node badge: warm orange (#f4a261).')

h3(doc, 'Presets (6)')
bullet(doc, '**None** — bypass, passthrough')
bullet(doc, '**Warm UGC** — subtle warmth + mild contrast lift, reads as "phone-shot but good"')
bullet(doc, '**Film Grain** — adds noise, slight desaturation, moodier')
bullet(doc, '**Golden Hour** — strong warmth, highlight compression, romantic')
bullet(doc, '**Clean Pop** — increased contrast + saturation, energetic')
bullet(doc, '**Moody** — cool shadows, crushed blacks, cinematic')

h3(doc, 'Sliders (9 — all granular)')
bullet(doc, '**Warmth** (−100 to +100) — color temperature via `colortemperature` filter (6500K baseline)')
bullet(doc, '**Tint** (−100 to +100) — Green ↔ Magenta balance via `colorbalance gm`')
bullet(doc, '**Exposure** (−2 to +2) — brightness via `eq brightness`')
bullet(doc, '**Contrast** (0.5 to 2.0) — via `eq contrast`, 1.0 = neutral')
bullet(doc, '**Saturation** (0 to 2.0) — via `eq saturation`')
bullet(doc, '**Highlight** (−100 to +100) — compression via curves (0.75 anchor)')
bullet(doc, '**Shadow** (−100 to +100) — lift via curves (0.25 anchor)')
bullet(doc, '**Grain** (0 to 50) — film grain via `noise c0s` filter')
bullet(doc, '**Sharpness** (0 to 5) — unsharp mask via `unsharp` filter')

h3(doc, 'Pipeline position')
para(doc,
     'Grade runs on any upstream video (UGC Video, Frame Sandwich, Remotion Compositor). Batch-processes '
     'sequentially (not parallel — avoids GPU/encoder contention). Outputs H.264 MP4 at CRF 18, '
     'audio copied without re-encode.', size=9)

h3(doc, 'What it cannot do today')
bullet(doc, 'No custom LUT file upload (`.cube` files). The filter supports it, the UI does not expose it.')
bullet(doc, 'No vignette, no blur, no chromatic aberration. Trivial to add — each is one FFmpeg filter.')
bullet(doc, 'No trim/cut/concat. It grades only; editing is a different node.')
bullet(doc, 'No GPU acceleration (uses libx264 software encoder). Would need NVENC/VAAPI path.')
bullet(doc, 'No preview before commit — you hit Generate and wait for the full render.')

# ── Node inventory ───────────────────────────────────────────────────────
h2(doc, 'Full node inventory (23 nodes, all implemented)')

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
hdr[0].text = 'Node'
hdr[1].text = 'Role'
hdr[2].text = 'Backed by'
for c in hdr:
    for run in c.paragraphs[0].runs:
        run.bold = True
    shade_cell(c, 'C9A227')

rows = [
    ('Character', 'Persona profile with pain points, hooks, psychology', 'Local state'),
    ('Ingredient', 'Pain point / hook atom, wired into generators', 'Local state'),
    ('Type', 'Script type + conversion level tags', 'Local state'),
    ('Group', 'Visual organizer (resizable container)', 'Local state'),
    ('Generator', 'Core script + production prompts', 'Anthropic'),
    ('Output', 'Displays voice / image / video / caption prompts', 'Local state'),
    ('UGC-Gen', 'UGC-style script generator', 'Anthropic'),
    ('Niche-Gen', 'Visual storytelling script (slide format)', 'Anthropic'),
    ('Clip Splitter', 'Script → 5s clip definitions with Kling prompts', 'Anthropic'),
    ('Avatar Frame', 'Local folder scanner for reference images', 'Local FS'),
    ('Clip Frames', 'Batch first-frame images per clip', 'kie.ai (Nano Banana)'),
    ('UGC Video', 'Batch Kling 3.0 from clip + frame pairs', 'kie.ai (Kling 3.0)'),
    ('16-GAMI Art', 'Batch origami art slides', 'kie.ai (Nano Banana)'),
    ('GAMI ARES', '3-agent origami prompt generation', 'kie.ai (Nano Banana)'),
    ('Title Card', 'Batch text-on-paper first frames', 'kie.ai (Nano Banana)'),
    ('Frame Sandwich', 'First + last frame → Kling 3.0 video', 'kie.ai (Kling 3.0)'),
    ('Carousel', 'Render branded slides with art zones', 'Local Python (render.py)'),
    ('Video Prompt', 'Batch motion prompt generator', 'Anthropic'),
    ('KIE Img2Vid', 'Batch img2vid (Kling 2.6, 3.0, MiniMax)', 'kie.ai'),
    ('KIE (Sora-2)', 'Text-to-video (Sora-2)', 'kie.ai'),
    ('Remotion Compositor', 'Video-in-slide compositor', 'Local Remotion CLI'),
    ('FFmpeg Grade', 'Color grade + film grain', 'Local FFmpeg'),
    ('Blotato', 'Multi-platform social posting', 'Blotato MCP'),
]
for r in rows:
    row = table.add_row().cells
    row[0].text = r[0]
    row[1].text = r[1]
    row[2].text = r[2]
    for run in row[0].paragraphs[0].runs:
        run.bold = True

# ── Realistic ideas ──────────────────────────────────────────────────────
h2(doc, 'Realistic 1–4 day builds (ranked by ROI for your POV-pivot direction)')

h3(doc, '1-day wins — tight scope, high impact')

para(doc, '① Subtitle Burn-In Node', bold=True, size=11)
bullet(doc, 'Why it matters now: POV content gets scrolled past on mute. Subtitles are the #1 retention lever on Reels/TikTok/Shorts.')
bullet(doc, 'What it does: Takes a graded video + transcript (already produced by the Scribe step), renders word-timed subtitles via FFmpeg `subtitles=` or `drawtext=` filter, burns them in.')
bullet(doc, 'Existing infra: Transcript already exists in `pipeline/shortform/edit/transcripts/`. FFmpeg server endpoint already handles batch video input. Just add a new endpoint `/api/ffmpeg/subtitle` + a node.')
bullet(doc, 'Time: 6–8 hours including style variants (TikTok bottom-center, karaoke word-highlight, full-caption).')

para(doc, '② Custom LUT Uploader for FFmpeg Grade', bold=True, size=11)
bullet(doc, 'Why it matters: Your POV footage needs a "signature look" to feel like a brand, not a phone dump. LUTs give you cinematic grading in one click.')
bullet(doc, 'What it does: Upload `.cube` files → server stores in `pipeline/luts/` → FFmpeg Grade node gets a LUT dropdown alongside presets.')
bullet(doc, 'Existing infra: FFmpeg already supports `lut3d=file.cube`. Upload endpoint pattern exists (`/api/upload-image`). Copy-paste + swap.')
bullet(doc, 'Time: 3–4 hours. Grab 5 free LUTs from Rocket Stock, test, ship.')

para(doc, '③ Notion → Script Feeder Node', bold=True, size=11)
bullet(doc, 'Why it matters: Closes the loop. You write teleprompter scripts in Notion → node pulls latest Draft/Recorded → feeds into Clip Splitter or Video Prompt. No more copy-paste.')
bullet(doc, 'What it does: Queries the Teleprompter Scripts database via Notion REST API, lets you pick a page by Status filter, exposes Topic + script body + Key Terms as output.')
bullet(doc, 'Existing infra: Notion API key already in `.env`. `fetchNotionKeyTerms` helper already written for the shortform CLI. Lift it into the Canvas.')
bullet(doc, 'Time: 6 hours. One node, one server endpoint wrapper.')

para(doc, '④ Vignette + Blur FFmpeg Sliders', bold=True, size=11)
bullet(doc, 'Why it matters: POV glasses have a slight fisheye look already. A gentle vignette pulls focus to the center (where your monitor usually is). Blur for depth-of-field fake.')
bullet(doc, 'What it does: Two new sliders in the Grade node → `vignette=PI/4:mode=NB` and `boxblur=radius`.')
bullet(doc, 'Time: 1–2 hours. Trivial.')

h3(doc, '2-day builds')

para(doc, '⑤ Hook-Frame Thumbnail Generator', bold=True, size=11)
bullet(doc, 'Why it matters: Reels retention drops off a cliff at 0.3s. A good hook-frame is a paused first frame with burned-in hook text.')
bullet(doc, 'What it does: Extracts first N frames, picks best one (sharpness + face detection if possible), overlays bold text, exports as .jpg + .mp4 with 0.5s freeze.')
bullet(doc, 'Existing infra: FFmpeg can extract frames and composite text. No new services.')
bullet(doc, 'Time: 10–14 hours including a few typographic variants.')

para(doc, '⑥ Grade Preset Library Node', bold=True, size=11)
bullet(doc, 'Why it matters: You land a grade you love (say, the look that pairs with your top-performing reel). Save it. Apply it to the next 20 videos without touching sliders.')
bullet(doc, 'What it does: Save slider state + LUT to JSON in `pipeline/grades/`. Loader node with dropdown of your saved grades.')
bullet(doc, 'Time: 8–10 hours.')

para(doc, '⑦ Audio Ducking Node', bold=True, size=11)
bullet(doc, 'Why it matters: When you add background music under your voice, the music has to duck under your speech automatically. Manual audio work kills throughput.')
bullet(doc, 'What it does: Takes two audio tracks (voice + music), applies FFmpeg `sidechaincompress` filter, returns ducked mix.')
bullet(doc, 'Existing infra: FFmpeg supports it natively. New endpoint, new node.')
bullet(doc, 'Time: 10 hours.')

para(doc, '⑧ Video Trimmer Node', bold=True, size=11)
bullet(doc, 'Why it matters: POV recordings start before you start talking. Trim the dead air programmatically.')
bullet(doc, 'What it does: Uses Scribe transcript to find first/last spoken word timestamp, trims +0.3s / −0.3s around.')
bullet(doc, 'Existing infra: Scribe transcript is already produced. FFmpeg trim is one filter.')
bullet(doc, 'Time: 8 hours.')

h3(doc, '3–4 day builds (ambitious but scoped)')

para(doc, '⑨ Vision-Powered Motion-Match Node', bold=True, size=11)
bullet(doc, 'Why it matters: Frame Sandwich currently takes one shared motion prompt for all pairs. Each pair deserves a tailored prompt — "paper unfolds from left" vs "envelope opens downward" — based on what the first and last frame actually show.')
bullet(doc, 'What it does: Pass first + last frame to Claude vision model → get a 40-60 word Kling motion prompt per pair → feeds Frame Sandwich.')
bullet(doc, 'Existing infra: Claude vision via Anthropic API. Frame Sandwich already accepts prompt input.')
bullet(doc, 'Time: 2–3 days (quality tuning is the long pole).')

para(doc, '⑩ In-Browser Teleprompter Mode', bold=True, size=11)
bullet(doc, 'Why it matters: For the teleprompter format A/B test. Web-based teleprompter reads your Notion page, highlights words as you speak (via Web Speech API), scrolls automatically. Much better than phone.')
bullet(doc, 'What it does: New Canvas panel (not a node — a sidebar tool). Pulls from Notion DB, shows clean large text, supports keyboard/pedal advance.')
bullet(doc, 'Time: 3 days including the speech-to-cursor sync.')

para(doc, '⑪ Reels/TikTok Analytics Pull + Benchmark Viewer', bold=True, size=11)
bullet(doc, 'Why it matters: Your top-performing reel is a benchmark. Every new post should be compared to it. A node that pulls Instagram Insights for a given post → saves to Notion → shows a ratio chart.')
bullet(doc, 'What it does: Instagram Graph API call (you already have the account). Saves metrics to a new Notion database "Post Benchmarks".')
bullet(doc, 'Time: 3–4 days (most of it is dealing with Meta auth).')

# ── NOT to build ─────────────────────────────────────────────────────────
h2(doc, 'What NOT to build — scope creep warnings')
bullet(doc, '**Cartesian AR rig** — you already tabled this, and correctly. Return when your baseline pipeline is producing consistent winners.')
bullet(doc, '**Full audio DAW / multi-track mixer** — DAWs are hard; use Descript or Audacity.')
bullet(doc, '**Cloud storage migration (S3/GCS)** — local filesystem works fine for you. Migrate when you collaborate with others, not before.')
bullet(doc, '**Multi-character scene generation** — requires a complete rewrite of character prompting and dialogue attribution. High effort, unclear ROI for solo operator.')
bullet(doc, '**Scheduled social posting** — Blotato already handles this. Don\'t rebuild what a $20/mo tool does well.')
bullet(doc, '**GPU encoding (NVENC)** — only matters if you are rendering dozens of videos per day. You are rendering a few. Libx264 is fine.')
bullet(doc, '**Real-time video preview during grading** — nice to have, weeks to build. Ship once your grade preset library is mature.')

# ── Strategic fit ────────────────────────────────────────────────────────
h2(doc, 'Strategic fit for your current direction')
para(doc,
     'Your top-performing reel told you two things: (a) POV glasses + real work on monitors = authentic authority, '
     '(b) save-worthy reference content converts elite (high save and follow rates). Every future '
     'build decision should pass the test: does this make it easier to ship save-worthy POV reference '
     'content at a higher rate? If no, defer.')
para(doc, 'Through that lens, the 1-day stack is the obvious priority:')
bullet(doc, 'Subtitles → mute-viewer retention (every winning Reel has them)')
bullet(doc, 'LUT uploader → brand consistency (your POV content needs a look)')
bullet(doc, 'Notion feeder → speed (no copy-paste between tools)')
para(doc,
     'These three together probably double your throughput from "recording to posted" '
     'in the next two weeks. The ambitious builds are real options — but ship the 1-day wins first, '
     'measure output cadence, then decide. ',
     italic=True)

h2(doc, 'Appendix: Extension surface — easy vs hard')
h3(doc, 'Easy (existing infra supports it)')
bullet(doc, 'FFmpeg filter additions — any new filter is a few lines in the filter chain builder')
bullet(doc, 'New kie.ai models — endpoint is model-agnostic, just add UI')
bullet(doc, 'Output node variants — OutputNode accepts arbitrary label/icon data')
bullet(doc, 'Motion style templates — Video Prompt styles are plain strings in a config object')
bullet(doc, 'Carousel themes — render.py accepts template strings')

h3(doc, 'Hard (requires new infra)')
bullet(doc, 'GPU encoding — needs new ffmpeg build + conditional args')
bullet(doc, 'Real-time preview — needs streaming endpoint + browser player')
bullet(doc, 'Cloud storage — needs storage adapter layer across 3+ endpoints')
bullet(doc, 'Advanced motion matching — needs vision model wiring + edge tracing logic')
bullet(doc, 'Multi-asset carousel layouts — render.py rewrite + UI changes')
bullet(doc, 'Multi-character scripts — system prompt overhaul, 200+ lines')

# ── Save ─────────────────────────────────────────────────────────────────
out_dir = Path('pipeline/docs')
out_dir.mkdir(parents=True, exist_ok=True)
out_path = out_dir / f'Breadstick_Canvas_Overview_{date.today().isoformat()}.docx'
doc.save(str(out_path))
print(f'Saved: {out_path}')
